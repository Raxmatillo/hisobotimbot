from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE

class MonthlyReportGenerator:
    def __init__(self):
        self.document = Document()
        self._setup_styles()

    def add_page_break(self):
        """Sahifa bo'linishini qo'shish"""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)  # Sahifa bo'linishini qo'shish

    def _setup_styles(self):
        """Hujjat stillarini sozlash"""
        # Sarlavha stili
        style = self.document.styles['Title']  # Mavjud 'Title' stilini olish
        style.font.bold = True
        style.font.size = Pt(14)
        style.font.name = 'Times New Roman'  # Shriftsiz nomini o'zgartirish
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Matn stili
        style = self.document.styles.add_style('NormalText', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(14)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justify qilib tekkizlash

    def create_header(self, region: str, district: str, school_number: str, teacher_name: str, position: str, month: str, year: str):
        """Hisobot sarlavhasini yaratish"""
        month_name = month.title()
        header_text = (
            f"\n\n\n\n\n\n{region.capitalize()} viloyati {district.capitalize()} MMT \nbo'limiga qarashli {school_number}"
            f"{position} {teacher_name}ning\n {year} o'quv yilining "
            f"{month_name} oyida \namalga oshirgan ishlari yuzasidan\n"
        )
    
        header = self.document.add_paragraph(header_text, style='Title')
        header.runs[0].font.size = Pt(22)
        # "HISOBOT" so'zini qo'shish
        hisobot_paragraph = self.document.add_paragraph("HISOBOT", style='Title')
        hisobot_paragraph.runs[0].font.size = Pt(26)  # Fontni kattalashtirish
        hisobot_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_paragraph()  # Bo'sh qator

    def add_activity(self, title: str, description: str, image_paths: list = None):
        """
        Faoliyat haqida ma'lumot va rasmlarni qo'shish
        :param title: Sarlavha
        :param description: Tavsif matni
        :param image_paths: Rasmlar ro'yxati (bir nechta rasm yo'li)
        """
        # Har bir yangi faoliyat uchun yangi sahifa qo'shish
        if title:
            self.add_page_break()  # Yangi sahifa

            # Sarlavha
            activity_title = self.document.add_paragraph(title, style='NormalText')
            activity_title.runs[0].bold = True

        # Tavsif (Justify qilib tekkizlash)
        if description:
            self.document.add_paragraph(description, style='NormalText')

        # Rasmlarni qo'shish
        if image_paths:
            # Rasm yo'li mavjud bo'lsa
            self.document.add_picture(image_paths, width=Inches(6))

        self.document.add_paragraph()  # Bo'sh qator

    def add_signature_block(self, director_name: str, teacher_name: str, position: str):
        """Imzo blokini qo'shish"""
        signature = self.document.add_paragraph(style='NormalText')
        signature.add_run(f"Maktab direktori: ").bold = True
        signature.add_run(director_name)

        teacher_sign = self.document.add_paragraph(style='NormalText')
        teacher_sign.add_run(f"Maktab {position}i: ").bold = True
        teacher_sign.add_run(teacher_name)

    def save(self, filename: str):
        """Hisobotni saqlash"""
        self.document.save(filename)



def generate_monthly_report(
    region: str,
    district: str,
    school_number: str,
    teacher_name: str,
    position: str,
    director_name: str,
    month: str,
    year: int,
    activities: list,
):
    """
    Oylik hisobot yaratish
    :param school_number: Maktab raqami
    :param teacher_name: O'qituvchi ismi
    :param position: Lavozimi
    :param director_name: Direktor ismi
    :param month: Oy raqami (1-12)
    :param year: Yil
    :param activities: Faoliyatlar ro'yxati [{'title': str, 'description': str, 'image': str}]
    :param output_file: Fayl nomi
    """
    report = MonthlyReportGenerator()
    
    # Sarlavha qo'shish
    report.create_header(
        region=region,
        district=district,
        school_number=school_number,
        teacher_name=teacher_name,
        position=position,
        month=month,
        year=year
    )
    
    # Faoliyatlarni qo'shish
    for activity in activities:
        report.add_activity(
            title=activity["title"],
            description=activity['description'],
            image_paths=activity['image']  # Rasmlar ro'yxati
        )

 
    # Imzolarni qo'shish
    report.add_signature_block(
        director_name=director_name,
        teacher_name=teacher_name,
        position=position
    )

    # Saqlash
    report.save(f'@hisobotimbot-{teacher_name}_{month}.docx')


# # Dasturni ishlatish namunasi
# if __name__ == "__main__":

#     acts = [
#         {
#             'title': "`Zo'ravonlikdan holi maktab` oyligi",
#             'description': "Lorem ipsum dolor, sit amet consectetur adipisicing elit. Asperiores ratione neque voluptatem nihil minima similique, quia rerum praesentium maxime beatae, voluptas repellendus vel labore incidunt nobis. Quis id doloremque temporibus illum provident quia nostrum omnis explicabo eius deserunt, voluptas natus dolorem laudantium, excv id doloremque temporibus illum provident quia nostrum omnis explica",
#             'images': ['images/image.jpg', 'images/tepaga.jpg', 'images/pppp.png']
#         },
#         {
#             'title': 'So\'rovnomalar o\'tkazish',
#             'description': 'Maktabimizning 7,8,9,10,11-sinf o\'quvchi qizlaridan so\'rovnomalar olindi va tahlil qilindi. Taqdimotlarda O‘zbekiston Respublikasining “Bolalarni zo‘ravonlikning barcha shakllaridan himoya qilish to‘g‘risida”,  “Pedagoglar orasida hissiy emotsional charchoq va kasbiy stress profilaktikasi” haqida tavsiya va tushunchalar berildi',
#             'images': ['images/qiz.png', 'images/Chaqqonlik.png']
#         }
#     ]
    
#     activities = []
    
#     i = 1
#     for activity in acts:
#         data = {}
#         data["title"] = activity["title"]
#         data["description"] = activity["description"]
#         images = activity["images"]
#         create_telegram_collage(images=images, output_path=f'12345_{i}.jpg')
#         data["image"] = f"12345_{i}.jpg"
#         i+=1
#         activities.append(data)

#     print(activities, '...........')
#     generate_monthly_report(
#         region='andijon',
#         district='shahrixon',
#         school_number=56,
#         teacher_name='Mirhamidova Durdona',
#         position='psixolog',
#         director_name='N.Maripov',
#         month='Dekabr',  # Dekabr
#         year=2024,
#         activities=activities,
#     )

"""
title - 60
description - 300
"""